from dotenv import load_dotenv
from langchain_community.document_loaders import Docx2txtLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
import docx
import zipfile
from xml.etree import ElementTree as ET
import os

load_dotenv()

# ── Create a sample Word doc for testing ──────────────────────────────────────
def create_sample_docx():
    doc = docx.Document()

    doc.add_heading("AI Engineer Interview Guide", level=1)
    doc.add_paragraph("This guide covers everything needed for AI engineer interviews.")

    doc.add_heading("RAG Systems", level=2)
    doc.add_paragraph("RAG combines retrieval with generation for accurate answers.")
    doc.add_paragraph("It grounds LLM outputs in real documents to reduce hallucinations.")

    doc.add_heading("Vector Databases", level=2)
    doc.add_paragraph("Vector databases store embeddings for semantic search.")
    doc.add_paragraph("Chroma is best for local development and prototyping.")

    doc.add_heading("Agents", level=2)
    doc.add_paragraph("Agents use LLMs to reason and act autonomously.")
    doc.add_paragraph("LangGraph is the best framework for building stateful agents.")

    # Add a table
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "Tool"
    headers[1].text = "Purpose"
    headers[2].text = "When to use"

    data = [
        ("Chroma",    "Vector storage",  "Local dev"),
        ("Pinecone",  "Cloud vector DB", "Production"),
        ("LangGraph", "Agent framework", "Stateful agents"),
    ]
    for i, (tool, purpose, when) in enumerate(data, start=1):
        row = table.rows[i].cells
        row[0].text = tool
        row[1].text = purpose
        row[2].text = when

    data_dir = os.path.join(os.path.dirname(__file__), "..", "data")
    os.makedirs(data_dir, exist_ok=True)
    path = os.path.join(data_dir, "sample.docx")
    doc.save(path)
    print(f"Sample Word doc created: {path}")
    return path


# ── Load Word doc preserving structure + comments ─────────────────────────────
def load_docx_with_structure(path: str) -> list[Document]:
    """
    Load a Word document preserving:
    - Heading hierarchy as metadata
    - Paragraph text with section context
    - Table content as structured text
    - Comments extracted via XML parsing
    """
    doc = docx.Document(path)
    documents = []
    current_h1 = ""
    current_h2 = ""
    current_h3 = ""

    # ── Paragraphs with heading metadata ──────────────────────────────────────
    for para in doc.paragraphs:
        if not para.text.strip():
            continue

        if para.style.name == "Heading 1":
            current_h1 = para.text
            current_h2 = ""
            current_h3 = ""
            continue
        elif para.style.name == "Heading 2":
            current_h2 = para.text
            current_h3 = ""
            continue
        elif para.style.name == "Heading 3":
            current_h3 = para.text
            continue

        documents.append(Document(
            page_content=para.text,
            metadata={
                "source": os.path.basename(path),
                "type":   "paragraph",
                "h1":     current_h1,
                "h2":     current_h2,
                "h3":     current_h3,
                "style":  para.style.name,
            }
        ))

    # ── Tables as structured text ──────────────────────────────────────────────
    for table_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells
            )
            rows.append(row_text)

        table_text = "\n".join(rows)
        documents.append(Document(
            page_content=table_text,
            metadata={
                "source":      os.path.basename(path),
                "type":        "table",
                "table_index": table_idx,
                "h1":          current_h1,
                "h2":          current_h2,
            }
        ))

    # ── Comments via XML ───────────────────────────────────────────────────────
    with zipfile.ZipFile(path, "r") as z:
        if "word/comments.xml" in z.namelist():
            with z.open("word/comments.xml") as f:
                tree = ET.parse(f)
                root = tree.getroot()
                ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

                for comment in root.findall(".//w:comment", ns):
                    author = comment.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author",
                        "Unknown"
                    )
                    date = comment.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date",
                        ""
                    )
                    texts = [
                        t.text for t in comment.findall(".//w:t", ns)
                        if t.text
                    ]
                    comment_text = " ".join(texts)

                    if comment_text.strip():
                        documents.append(Document(
                            page_content=comment_text,
                            metadata={
                                "source": os.path.basename(path),
                                "type":   "comment",
                                "author": author,
                                "date":   date,
                            }
                        ))

    return documents


# ── Main ──────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    doc_path = create_sample_docx()

    # Method 1: Basic LangChain loader
    print("\n" + "=" * 60)
    print("Method 1 — Docx2txtLoader (basic, loses structure)")
    print("=" * 60)
    basic_loader = Docx2txtLoader(doc_path)
    basic_docs = basic_loader.load()
    print(f"Loaded {len(basic_docs)} document")
    print(f"Content preview: {basic_docs[0].page_content[:200]}...")

    # Method 2: python-docx with full structure
    print("\n" + "=" * 60)
    print("Method 2 — python-docx (full structure + comments)")
    print("=" * 60)
    structured_docs = load_docx_with_structure(doc_path)
    print(f"Total chunks extracted: {len(structured_docs)}")
    for doc in structured_docs:
        print(f"  [{doc.metadata['type']}] "
              f"h1='{doc.metadata.get('h1', '')}' "
              f"h2='{doc.metadata.get('h2', '')}' → "
              f"{doc.page_content[:70]}...")

    # Store in Chroma
    embeddings = HuggingFaceEmbeddings(
        model_name="all-MiniLM-L6-v2",
        model_kwargs={"device": "cpu"}
    )
    persist_dir = os.path.join(os.path.dirname(__file__), "..", "chroma_db")
    store = Chroma.from_documents(
        structured_docs,
        embeddings,
        persist_directory=persist_dir,
        collection_name="word_docs"
    )

    query = "Which tool should I use for production?"
    print(f"\nQuery: '{query}'")

    print("\nFilter type='table':")
    table_results = store.similarity_search(query, k=2, filter={"type": "table"})
    for doc in table_results:
        print(f"  {doc.page_content}")

    print("\nFilter type='comment':")
    comment_results = store.similarity_search(query, k=2, filter={"type": "comment"})
    if comment_results:
        for doc in comment_results:
            print(f"  [by {doc.metadata['author']}] {doc.page_content}")
    else:
        print("  No comments in sample doc — real docs will populate this")